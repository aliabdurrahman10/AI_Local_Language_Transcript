# ============================================
# Script transkripsi batch audio dengan Google Gemini
# Fokus:
# 1. Membaca semua file audio dalam folder Audio
# 2. Upload file ke Files API Gemini
# 3. Meminta transkrip sedetail mungkin
# 4. Menyimpan hasil ke TXT, DOCX, PDF, dan JSON
# ============================================

import json
import mimetypes
import os
import sys
import time
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from google import genai
from google.genai import types
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from reportlab.pdfgen import canvas


# ============================================
# Bagian 1 - Pengaturan dasar
# Ubah sesuai folder di komputer
# ============================================

BASE_DIR = Path(r"C:\Users\Ali\Documents\Repository Github\Transkrip")
INPUT_DIR = BASE_DIR / "Audio"
OUTPUT_DIR = BASE_DIR / "Hasil"

# Model terbaik untuk kualitas
# Kalau nanti terasa terlalu mahal atau lambat, bisa ganti ke "gemini-2.5-flash"
MODEL_NAME = "gemini-2.5-flash"

# Format audio yang didukung Gemini
SUPPORTED_EXTENSIONS = {".mp3", ".aac", ".wav", ".m4a", ".ogg", ".flac", ".aiff"}

# Waktu tunggu maksimal saat file diproses server
MAX_WAIT_SECONDS = 60 * 20
POLL_INTERVAL_SECONDS = 3

# Simpan hasil mentah respons API juga
SAVE_JSON_DEBUG = True


# ============================================
# Bagian 2 - API key dan client
# Script akan mencoba GEMINI_API_KEY dulu,
# lalu GOOGLE_API_KEY kalau variabel pertama tidak ada
# ============================================

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise RuntimeError(
        "API key tidak ditemukan. Isi file .env dengan GEMINI_API_KEY=... "
        "atau GOOGLE_API_KEY=..."
    )

client = genai.Client(api_key=api_key)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================
# Bagian 3 - Fungsi bantu
# ============================================

def safe_filename(name: str) -> str:
    bad_chars = r'<>:"/\|?*'
    for ch in bad_chars:
        name = name.replace(ch, "_")
    return name


def guess_mime_type(path: Path) -> str:
    mapping = {
        ".mp3": "audio/mp3",
        ".aac": "audio/aac",
        ".wav": "audio/wav",
        ".m4a": "audio/mp4",
        ".ogg": "audio/ogg",
        ".flac": "audio/flac",
        ".aiff": "audio/aiff",
    }
    guessed = mapping.get(path.suffix.lower())
    if guessed:
        return guessed

    mime, _ = mimetypes.guess_type(str(path))
    if mime:
        return mime

    return "application/octet-stream"


def save_txt(path: Path, text: str):
    path.write_text(text, encoding="utf-8")


def save_json(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def save_docx(path: Path, title: str, text: str):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(path))


def save_pdf(path: Path, title: str, text: str):
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 40
    y = height - margin

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, title)
    y -= 24

    c.setFont("Helvetica", 10)

    for para in text.split("\n"):
        wrapped_lines = simpleSplit(para, "Helvetica", 10, width - (2 * margin))
        if not wrapped_lines:
            wrapped_lines = [""]

        for line in wrapped_lines:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - margin
            c.drawString(margin, y, line)
            y -= 14
        y -= 4

    c.save()


def extract_response_text(response) -> str:
    text = getattr(response, "text", None)
    if text:
        return text.strip()

    try:
        if hasattr(response, "candidates") and response.candidates:
            parts = response.candidates[0].content.parts
            joined = []
            for part in parts:
                part_text = getattr(part, "text", None)
                if part_text:
                    joined.append(part_text)
            if joined:
                return "\n".join(joined).strip()
    except Exception:
        pass

    return ""


def upload_audio_file(audio_path: Path):
    mime_type = guess_mime_type(audio_path)
    return client.files.upload(
        file=str(audio_path),
        config={"mime_type": mime_type}
    )


def wait_until_ready(uploaded_file):
    started = time.time()
    current = uploaded_file

    while True:
        state_obj = getattr(current, "state", None)
        state_name = getattr(state_obj, "name", None) if state_obj else None

        if not state_name:
            return current

        if state_name.upper() in {"ACTIVE", "READY"}:
            return current

        if state_name.upper() == "FAILED":
            raise RuntimeError("File gagal diproses oleh server Gemini.")

        waited = time.time() - started
        if waited > MAX_WAIT_SECONDS:
            raise TimeoutError("Waktu tunggu file di server Gemini terlalu lama.")

        time.sleep(POLL_INTERVAL_SECONDS)
        current = client.files.get(name=current.name)


def build_prompt(audio_name: str) -> str:
    return f"""
Tolong transkripkan file audio berikut dengan kualitas setinggi mungkin.

Tujuan:
- menghasilkan transkrip VERBATIM
- sedekat mungkin dengan ucapan asli
- mempertahankan campuran Bahasa Indonesia dan dialek Bengkulu
- jangan merapikan isi secara berlebihan
- jangan meringkas
- jangan menghilangkan kata pengisi bila terdengar jelas

Konteks:
- Ini wawancara lapangan tentang nelayan, PLTU, tangkapan ikan, kerusakan karang, limbah air panas, polusi batu bara, biaya solar, dan kondisi Teluk Sepang.
- Pewawancara bisa jadi mahasiswa/peneliti.
- Penutur dapat menggunakan Bahasa Indonesia campur dialek Bengkulu.

Aturan penulisan:
1. Gunakan format transkrip dengan timestamp.
2. Tulis timestamp setiap pergantian pembicara atau saat topik berpindah jelas.
3. Format yang dipakai:
   [HH:MM:SS] Pembicara 1: isi ucapan
4. Kalau identitas pembicara tidak jelas, gunakan label netral seperti Pembicara 1, Pembicara 2, dst.
5. Kalau ada suara latar penting, tulis dalam kurung siku, misalnya [suara ombak], [suara mesin], [tertawa], [tidak jelas].
6. Kalau ada bagian yang tidak terdengar jelas, tulis [tidak jelas] dan jangan mengarang.
7. Jangan membuat kesimpulan. Hanya transkrip.
8. Utamakan ketelitian kata dan urutan ujaran.
9. Jangan mengubah dialek Bengkulu menjadi bahasa baku.
10. Jangan keluarkan penjelasan tambahan di luar transkrip.

Nama file:
{audio_name}
""".strip()


def generate_transcript(uploaded_file, audio_name: str):
    prompt = build_prompt(audio_name)

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=[prompt, uploaded_file],
        config=types.GenerateContentConfig(
            system_instruction=(
                "Anda adalah transcriber profesional untuk wawancara lapangan. "
                "Anda harus menghasilkan transkrip verbatim yang jujur, teliti, "
                "dan tidak mengarang bagian yang tidak terdengar."
            ),
            temperature=0.1,
            top_p=0.95,
            candidate_count=1,
            max_output_tokens=65536,
        ),
    )

    return response


def process_one_file(audio_path: Path):
    uploaded_file = None

    try:
        print(f"\nMemproses file: {audio_path.name}")
        print("  [1/4] Upload file ke Gemini...")
        uploaded_file = upload_audio_file(audio_path)

        print("  [2/4] Menunggu file siap diproses...")
        uploaded_file = wait_until_ready(uploaded_file)

        print("  [3/4] Membuat transkrip...")
        response = generate_transcript(uploaded_file, audio_path.name)
        text = extract_response_text(response)

        if not text.strip():
            raise RuntimeError("Respons Gemini kosong.")

        base_name = safe_filename(audio_path.stem)
        title = f"Transkrip - {audio_path.name}"

        out_txt = OUTPUT_DIR / f"{base_name}_GEMINI.txt"
        out_docx = OUTPUT_DIR / f"{base_name}_GEMINI.docx"
        out_pdf = OUTPUT_DIR / f"{base_name}_GEMINI.pdf"
        out_json = OUTPUT_DIR / f"{base_name}_GEMINI.json"

        print("  [4/4] Menyimpan hasil...")
        save_txt(out_txt, text)
        save_docx(out_docx, title, text)
        save_pdf(out_pdf, title, text)

        if SAVE_JSON_DEBUG:
            try:
                debug_payload = response.model_dump()
            except Exception:
                debug_payload = {"text": text}
            save_json(out_json, debug_payload)

        print(f"  [SUKSES] {audio_path.name}")
        return {
            "file": audio_path.name,
            "status": "success",
            "text_file": str(out_txt),
            "docx_file": str(out_docx),
            "pdf_file": str(out_pdf),
        }

    except Exception as e:
        print(f"  [GAGAL] {audio_path.name}: {e}")
        return {
            "file": audio_path.name,
            "status": "failed",
            "error": str(e),
        }

    finally:
        if uploaded_file is not None:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass


# ============================================
# Bagian 4 - Proses utama
# ============================================

def main():
    files = sorted([
        p for p in INPUT_DIR.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    ])

    if not files:
        print(f"Tidak ada file audio di folder: {INPUT_DIR}")
        sys.exit(1)

    print(f"Model     : {MODEL_NAME}")
    print(f"Folder    : {INPUT_DIR}")
    print(f"Jumlah    : {len(files)} file")

    results = []

    for audio_path in files:
        result = process_one_file(audio_path)
        results.append(result)

    success_count = sum(1 for r in results if r["status"] == "success")
    failed_count = sum(1 for r in results if r["status"] == "failed")

    summary_path = OUTPUT_DIR / "ringkasan_proses_gemini.json"
    save_json(summary_path, results)

    print("\nProses selesai.")
    print(f"Berhasil : {success_count}")
    print(f"Gagal    : {failed_count}")
    print(f"Ringkasan: {summary_path}")
    print(f"Hasil    : {OUTPUT_DIR}")

    if success_count == 0:
        sys.exit(2)


if __name__ == "__main__":
    main()
