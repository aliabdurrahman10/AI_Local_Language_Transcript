# ============================================
# Script transkripsi batch audio dengan OpenAI
# Versi revisi
#
# Perbaikan utama:
# 1. File .aac tidak dikirim langsung ke API
#    tetapi dikonversi dulu ke format upload yang aman.
# 2. Ada pengecekan ffmpeg sebelum konversi/potong audio.
# 3. Bisa pilih mode:
#    - best_text  : fokus akurasi teks
#    - detailed   : fokus speaker + timestamp
# 4. Menentukan language='id' untuk bantu akurasi.
# 5. Menyimpan log file yang gagal diproses.
# 6. Kalau semua file gagal, script keluar dengan status gagal.
# ============================================

import os
import sys
import math
import json
import shutil
import tempfile
from pathlib import Path

from dotenv import load_dotenv
from openai import OpenAI
from pydub import AudioSegment
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


# ============================================
# Bagian 1 - Pengaturan dasar
# ============================================

BASE_DIR = Path(r"C:\Users\Ali\Documents\Repository Github\Transkrip")
INPUT_DIR = BASE_DIR / "Audio"
OUTPUT_DIR = BASE_DIR / "Hasil"

# Pilihan mode:
# - "best_text" : teks paling aman untuk dibaca
# - "detailed"  : speaker + timestamp
MODE = "detailed"

TEXT_MODEL = "gpt-4o-transcribe"
DETAILED_MODEL = "gpt-4o-transcribe-diarize"

# Bahasa input untuk bantu akurasi
LANGUAGE = "id"

# Petunjuk singkat untuk model teks
TRANSCRIPT_PROMPT = (
    "Transkripsikan percakapan apa adanya dalam bahasa Indonesia. "
    "Audio bisa memuat kata serapan lokal dan dialek Bengkulu. "
    "Jangan merangkum. Pertahankan ucapan tidak baku sejauh mungkin."
)

# Format input lokal yang diterima script
SUPPORTED_INPUT_EXTENSIONS = {
    ".aac", ".flac", ".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".ogg", ".wav", ".webm"
}

# Format upload yang aman untuk API
SUPPORTED_UPLOAD_EXTENSIONS = {
    ".flac", ".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".ogg", ".wav", ".webm"
}

MAX_FILE_SIZE_MB = 25
SAFE_MAX_FILE_SIZE_MB = 24

# Format hasil konversi lokal
CONVERT_EXPORT_FORMAT = "mp3"
CONVERT_BITRATE = "96k"
TARGET_SAMPLE_RATE = 16000
TARGET_CHANNELS = 1


# ============================================
# Bagian 2 - Persiapan awal
# ============================================

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise RuntimeError("OPENAI_API_KEY belum ditemukan. Isi dulu file .env")

client = OpenAI(api_key=api_key)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================
# Bagian 3 - Fungsi bantu umum
# ============================================


def safe_filename(name: str) -> str:
    bad_chars = r'<>:"/\\|?*'
    for ch in bad_chars:
        name = name.replace(ch, "_")
    return name



def file_size_mb(file_path: Path) -> float:
    return file_path.stat().st_size / (1024 * 1024)



def format_seconds(seconds):
    if seconds is None:
        return "00:00:00"
    total = int(float(seconds))
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"



def save_txt(path: Path, text: str):
    path.write_text(text, encoding="utf-8")



def save_json(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")



def wrap_line_for_pdf(text, max_width, font_name="Helvetica", font_size=10):
    words = text.split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        trial = current + " " + word
        if stringWidth(trial, font_name, font_size) <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines



def save_docx(path: Path, title: str, text: str):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(path))



def save_pdf(path: Path, title: str, text: str):
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 40
    y = height - margin

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, title)
    y -= 24

    c.setFont("Helvetica", 10)
    max_width = width - 2 * margin

    for para in text.split("\n"):
        lines = wrap_line_for_pdf(para, max_width)
        for line in lines:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - margin
            c.drawString(margin, y, line)
            y -= 14
        y -= 4

    c.save()



def ffmpeg_available() -> bool:
    return shutil.which("ffmpeg") is not None or shutil.which("ffmpeg.exe") is not None



def require_ffmpeg():
    if not ffmpeg_available():
        raise RuntimeError(
            "ffmpeg belum ditemukan. Install ffmpeg dan pastikan ffmpeg bisa dipanggil dari terminal."
        )


# ============================================
# Bagian 4 - Persiapan audio sebelum upload
# ============================================


def estimate_chunk_count(file_mb: float) -> int:
    return max(1, math.ceil(file_mb / SAFE_MAX_FILE_SIZE_MB) + 1)



def needs_local_conversion(audio_path: Path) -> bool:
    ext = audio_path.suffix.lower()
    if ext not in SUPPORTED_UPLOAD_EXTENSIONS:
        return True
    return False



def normalize_audio(audio: AudioSegment) -> AudioSegment:
    return audio.set_frame_rate(TARGET_SAMPLE_RATE).set_channels(TARGET_CHANNELS)



def prepare_audio_for_upload(audio_path: Path):
    """
    Menghasilkan daftar file siap upload ke API.
    Kalau file sudah aman, dipakai langsung.
    Kalau format tidak aman atau perlu dipotong, file dikonversi dulu.
    """
    size_mb = file_size_mb(audio_path)
    ext = audio_path.suffix.lower()

    # File aman: format didukung API dan ukuran aman
    if ext in SUPPORTED_UPLOAD_EXTENSIONS and size_mb <= SAFE_MAX_FILE_SIZE_MB:
        return [{
            "chunk_index": 1,
            "start_ms": 0,
            "end_ms": None,
            "path": audio_path,
            "temporary": False,
        }]

    require_ffmpeg()

    audio = AudioSegment.from_file(str(audio_path))
    audio = normalize_audio(audio)

    # Kalau cuma formatnya yang perlu diubah, buat satu file konversi saja
    if size_mb <= SAFE_MAX_FILE_SIZE_MB and needs_local_conversion(audio_path):
        temp_dir = Path(tempfile.mkdtemp(prefix="audio_upload_ready_"))
        out_name = f"{safe_filename(audio_path.stem)}_upload.{CONVERT_EXPORT_FORMAT}"
        out_path = temp_dir / out_name
        audio.export(str(out_path), format=CONVERT_EXPORT_FORMAT, bitrate=CONVERT_BITRATE)

        if file_size_mb(out_path) > MAX_FILE_SIZE_MB:
            raise ValueError(
                f"File hasil konversi {out_path.name} masih terlalu besar ({file_size_mb(out_path):.2f} MB)."
            )

        return [{
            "chunk_index": 1,
            "start_ms": 0,
            "end_ms": None,
            "path": out_path,
            "temporary": True,
        }]

    # Kalau perlu dipotong
    chunk_count = estimate_chunk_count(size_mb)
    total_ms = len(audio)
    chunk_length_ms = math.ceil(total_ms / chunk_count)

    chunks = []
    temp_dir = Path(tempfile.mkdtemp(prefix="audio_chunks_"))

    for i in range(chunk_count):
        start_ms = i * chunk_length_ms
        end_ms = min((i + 1) * chunk_length_ms, total_ms)
        piece = audio[start_ms:end_ms]

        chunk_name = f"{safe_filename(audio_path.stem)}_bagian_{i+1}.{CONVERT_EXPORT_FORMAT}"
        chunk_path = temp_dir / chunk_name

        piece.export(str(chunk_path), format=CONVERT_EXPORT_FORMAT, bitrate=CONVERT_BITRATE)

        chunk_mb = file_size_mb(chunk_path)
        if chunk_mb > MAX_FILE_SIZE_MB:
            raise ValueError(
                f"Potongan {chunk_path.name} masih terlalu besar ({chunk_mb:.2f} MB). "
                f"Turunkan bitrate atau perbanyak jumlah potongan."
            )

        chunks.append({
            "chunk_index": i + 1,
            "start_ms": start_ms,
            "end_ms": end_ms,
            "path": chunk_path,
            "temporary": True,
        })

    return chunks


# ============================================
# Bagian 5 - Panggilan API transkripsi
# ============================================


def transcribe_one_file(audio_file: Path):
    with open(audio_file, "rb") as f:
        kwargs = {
            "file": f,
            "language": LANGUAGE,
            "temperature": 0,
        }

        if MODE == "detailed":
            kwargs["model"] = DETAILED_MODEL
            kwargs["response_format"] = "diarized_json"
            kwargs["chunking_strategy"] = "auto"
        else:
            kwargs["model"] = TEXT_MODEL
            kwargs["response_format"] = "json"
            kwargs["prompt"] = TRANSCRIPT_PROMPT
            kwargs["include"] = ["logprobs"]

        result = client.audio.transcriptions.create(**kwargs)

    if hasattr(result, "model_dump"):
        data = result.model_dump()
    elif isinstance(result, dict):
        data = result
    else:
        data = json.loads(json.dumps(result, default=str))

    return data


# ============================================
# Bagian 6 - Menyusun hasil
# ============================================


def shift_segments(segments, offset_seconds):
    shifted = []

    for seg in segments:
        new_seg = dict(seg)

        if new_seg.get("start") is not None:
            new_seg["start"] = float(new_seg["start"]) + offset_seconds

        if new_seg.get("end") is not None:
            new_seg["end"] = float(new_seg["end"]) + offset_seconds

        shifted.append(new_seg)

    return shifted



def build_text_from_segments(segments, fallback_text=""):
    if not segments:
        return fallback_text.strip()

    lines = []
    for seg in segments:
        start = format_seconds(seg.get("start"))
        end = format_seconds(seg.get("end"))
        speaker = seg.get("speaker", "Pembicara")
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        lines.append(f"[{start} - {end}] {speaker}: {text}")

    return "\n".join(lines).strip()



def transcribe_full_audio(audio_path: Path):
    chunks = prepare_audio_for_upload(audio_path)

    full_text_parts = []
    full_segments = []
    raw_parts = []

    for chunk_info in chunks:
        chunk_path = chunk_info["path"]
        start_ms = chunk_info["start_ms"]
        offset_seconds = start_ms / 1000.0

        print(f"  Memproses potongan: {chunk_path.name}")
        data = transcribe_one_file(chunk_path)

        raw_parts.append({
            "chunk_index": chunk_info["chunk_index"],
            "start_ms": start_ms,
            "end_ms": chunk_info["end_ms"],
            "source_file": str(chunk_path),
            "result": data,
        })

        piece_text = (data.get("text") or "").strip()
        if piece_text:
            full_text_parts.append(piece_text)

        segments = data.get("segments", []) or []
        if segments:
            shifted = shift_segments(segments, offset_seconds)
            full_segments.extend(shifted)

    merged_text = "\n".join(full_text_parts).strip()
    transcript_text = build_text_from_segments(full_segments, fallback_text=merged_text)

    if not transcript_text.strip():
        transcript_text = "[Transkrip kosong atau tidak terbaca]"

    merged_data = {
        "source_file": str(audio_path),
        "mode": MODE,
        "text": merged_text,
        "segments": full_segments,
        "parts": raw_parts,
    }

    return merged_data, transcript_text


# ============================================
# Bagian 7 - Proses utama
# ============================================


def main():
    audio_files = sorted([
        p for p in INPUT_DIR.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_INPUT_EXTENSIONS
    ])

    if not audio_files:
        print(f"Tidak ada file audio di folder: {INPUT_DIR}")
        return

    combined_doc = Document()
    combined_doc.add_heading("Transkrip Gabungan", level=1)

    combined_text_blocks = []
    failed_files = []
    success_count = 0

    print(f"Mode: {MODE}")
    print(f"Folder input: {INPUT_DIR}")
    print(f"Jumlah file audio: {len(audio_files)}")

    for audio_path in audio_files:
        print(f"\nMemproses file: {audio_path.name}")

        try:
            merged_data, transcript_text = transcribe_full_audio(audio_path)
        except Exception as e:
            msg = f"{audio_path.name}: {e}"
            print(f"Gagal memproses {msg}")
            failed_files.append(msg)
            continue

        base_name = safe_filename(audio_path.stem)
        out_txt = OUTPUT_DIR / f"{base_name}.txt"
        out_json = OUTPUT_DIR / f"{base_name}.json"
        out_docx = OUTPUT_DIR / f"{base_name}.docx"
        out_pdf = OUTPUT_DIR / f"{base_name}.pdf"

        save_txt(out_txt, transcript_text)
        save_json(out_json, merged_data)
        save_docx(out_docx, f"Transkrip - {audio_path.name}", transcript_text)
        save_pdf(out_pdf, f"Transkrip - {audio_path.name}", transcript_text)

        combined_doc.add_heading(audio_path.name, level=2)
        for line in transcript_text.split("\n"):
            combined_doc.add_paragraph(line)

        combined_text_blocks.append(f"===== {audio_path.name} =====\n{transcript_text}\n")
        success_count += 1

    combined_txt = OUTPUT_DIR / "GABUNGAN_Transkrip.txt"
    combined_json = OUTPUT_DIR / "GABUNGAN_Transkrip.json"
    combined_docx = OUTPUT_DIR / "GABUNGAN_Transkrip.docx"
    combined_pdf = OUTPUT_DIR / "GABUNGAN_Transkrip.pdf"
    failed_log = OUTPUT_DIR / "GAGAL_per_file.txt"

    combined_text = "\n".join(combined_text_blocks)

    save_txt(combined_txt, combined_text)
    save_json(combined_json, {"files_combined_text": combined_text_blocks})
    combined_doc.save(str(combined_docx))
    save_pdf(combined_pdf, "GABUNGAN_Transkrip", combined_text)

    if failed_files:
        save_txt(failed_log, "\n".join(failed_files))

    print("\nProses selesai.")
    print(f"Berhasil: {success_count}")
    print(f"Gagal   : {len(failed_files)}")
    print(f"Hasil tersimpan di folder: {OUTPUT_DIR}")

    if success_count == 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
